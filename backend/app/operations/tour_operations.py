from typing import List
from database.database import get_connection
from schemas import models
import os
from dotenv import load_dotenv
import requests
from bs4 import BeautifulSoup
import urllib.parse
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.shared import OxmlElement, qn
from io import BytesIO

load_dotenv()
API_URL = os.getenv('API_URL', 'http://127.0.0.1:8002/api/v1/search_location')

def save_tour_categories(tour_id: int, categories: List[str]):
    """Save tour categories to database"""
    conn = get_connection()
    cursor = conn.cursor()
    
    try:
        # Delete existing categories
        cursor.execute("DELETE FROM tour_categories WHERE tour_id = ?", (tour_id,))
        
        # Insert new categories
        for category in categories:
            cursor.execute('''
            INSERT INTO tour_categories (tour_id, category)
            VALUES (?, ?)
            ''', (tour_id, category))
        
        conn.commit()
        return {"status": "success", "message": "Категории успешно сохранены"}
    except Exception as e:
        conn.rollback()
        return {"status": "error", "message": str(e)}
    finally:
        conn.close()

def get_tour_categories(tour_id: int) -> List[str]:
    """Get tour categories from database"""
    conn = get_connection()
    cursor = conn.cursor()
    
    try:
        # cursor.execute("SELECT category FROM tour_categories WHERE tour_id = ?", (tour_id,))
        # return [row[0] for row in cursor.fetchall()]
        return []
    finally:
        conn.close()

def save_tour_to_db(tour_data, url=None):
    """Save tour data to database"""
    conn = get_connection()
    cursor = conn.cursor()
    
    try:
        # Insert tour
        cursor.execute('''
        INSERT INTO tours (title, date_start, date_end, location, rating, relevance, url)
        VALUES (?, ?, ?, ?, ?, ?, ?)
        ''', (
            tour_data.title,
            tour_data.date[0] if hasattr(tour_data, 'date') and len(tour_data.date) > 0 else None,
            tour_data.date[1] if hasattr(tour_data, 'date') and len(tour_data.date) > 1 else None,
            tour_data.location,
            tour_data.rating,
            tour_data.relevance,
            url
        ))
        
        tour_id = cursor.lastrowid
        
        # Insert places
        for place in tour_data.places:
            cursor.execute('''
            INSERT INTO places (tour_id, name, location, rating, date_start, date_end, 
                              description, photo, mapgeo_x, mapgeo_y)
            VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
            ''', (
                tour_id,
                place.name,
                place.location,
                place.rating,
                getattr(place, 'date_start', None) if hasattr(place, 'date_start') else None,
                getattr(place, 'date_end', None) if hasattr(place, 'date_end') else None,
                place.description,
                place.photo,
                place.mapgeo[0],
                place.mapgeo[1]
            ))
        
        
        
        conn.commit()
        return tour_id
    except Exception as e:
        conn.rollback()
        raise e
    finally:
        conn.close()

def get_first_google_image_url(place_name):
    """Get first Google image URL for place name"""
    try:
        search_query = f"{place_name} достопримечательность"
        url = f"https://www.google.com/search?q={urllib.parse.quote(search_query)}&tbm=isch"
        
        headers = {
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36'
        }
        
        response = requests.get(url, headers=headers)
        soup = BeautifulSoup(response.content, 'html.parser')
        
        # Find first image
        img_tags = soup.find_all('img')
        for img in img_tags:
            src = img.get('src')
            if src and src.startswith('http'):
                return src
                
        return "https://via.placeholder.com/150"
    except:
        return "https://via.placeholder.com/150"

def get_tour_by_id(tour_id: int):
    """Get tour data by ID"""
    print(f"Getting tour by ID: {tour_id}")
    conn = get_connection()
    cursor = conn.cursor()
    
    try:
        # Get tour data
        cursor.execute('''
        SELECT title, date_start, date_end, location, rating, relevance, url
        FROM tours
        WHERE tour_id = ?
        ''', (tour_id,))
        
        tour_row = cursor.fetchone()
        print(f"Tour row found: {tour_row}")
        if not tour_row:
            print(f"No tour found for ID: {tour_id}")
            return None
        
        # Get places
        cursor.execute('''
        SELECT name, location, rating, date_start, date_end, description, photo, mapgeo_x, mapgeo_y
        FROM places
        WHERE tour_id = ?
        ''', (tour_id,))
        
        places_rows = cursor.fetchall()
        print(f"Found {len(places_rows)} places for tour {tour_id}")
        
        places = []
        
        for place_row in places_rows:
            place_row = [0] + list(place_row)
            photo_url = place_row[7]
            if not photo_url or photo_url == "https://via.placeholder.com/150":
                photo_url = get_first_google_image_url(place_row[1])
            places.append(models.Places(
                id_place=place_row[0],
                name=place_row[1],
                location=place_row[2],
                rating=place_row[3],
                date=f"{place_row[4]} - {place_row[5]}" if place_row[4] and place_row[5] else str(place_row[4] or place_row[5]),
                description=place_row[6],
                photo=photo_url,
                mapgeo=[place_row[8], place_row[9]]
            ))
        
        # Get categories
        categories = get_tour_categories(tour_id)
        
        tour = models.Tour(
            tour_id=tour_id,
            title=tour_row[0],
            date=[tour_row[1] or '', tour_row[2] or ''],
            location=tour_row[3],
            rating=tour_row[4],
            relevance=tour_row[5],
            url=tour_row[6],
            places=places,
            categories=categories,
            description="Увлекательный тур для всей семьи!"
        )
        print(f"Successfully created tour object for ID: {tour_id}")
        return tour
    except Exception as e:
        print(f"Error in get_tour_by_id for tour {tour_id}: {str(e)}")
        raise e
    finally:
        conn.close()

def get_popular_tours():
    """Get popular tours from database"""
    conn = get_connection()
    cursor = conn.cursor()
    
    try:
        cursor.execute('''
        SELECT tour_id, title, date_start, date_end, location, rating, relevance, url
        FROM tours
        ORDER BY rating DESC, relevance DESC
        LIMIT 10
        ''')
        
        tours = []
        for row in cursor.fetchall():
            tour_id = row[0]
            tour = get_tour_by_id(tour_id)
            if tour:
                tours.append(tour)
        
        return tours
    finally:
        conn.close()

def generate_tour_docx(tour_id: int) -> bytes:
    """Generate Word document for tour by ID"""
    # Get tour data
    tour = get_tour_by_id(tour_id)
    if not tour:
        raise ValueError("Tour not found")
    
    # Create document
    doc = Document()
    
    # Add title
    title = doc.add_heading(f'Тур: {tour.title}', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add tour information section
    doc.add_heading('Информация о туре', level=1)
    
    # Create tour info table
    tour_table = doc.add_table(rows=6, cols=2)
    tour_table.style = 'Table Grid'
    tour_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Fill tour info table
    tour_data = [
        ('Название:', tour.title),
        ('Локация:', tour.location),
        ('Дата начала:', tour.date[0] if tour.date and len(tour.date) > 0 else 'Не указана'),
        ('Дата окончания:', tour.date[1] if tour.date and len(tour.date) > 1 else 'Не указана'),
        ('Рейтинг:', str(tour.rating)),
        ('Описание:', tour.description or 'Описание отсутствует')
    ]
    
    for i, (key, value) in enumerate(tour_data):
        tour_table.cell(i, 0).text = key
        tour_table.cell(i, 1).text = value
    
    # Style the table header
    for cell in tour_table.rows[0].cells:
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].runs[0].font.size = Pt(12)
    
    # Add places section
    if tour.places:
        doc.add_heading('Места для посещения', level=1)
        
        for i, place in enumerate(tour.places, 1):
            # Add place name
            place_heading = doc.add_heading(f'{i}. {place.name}', level=2)
            
            # Create place info table
            place_table = doc.add_table(rows=4, cols=2)
            place_table.style = 'Table Grid'
            place_table.alignment = WD_TABLE_ALIGNMENT.CENTER
            
            # Fill place info table
            place_data = [
                ('Категория:', place.description or 'Не указана'),
                ('Рейтинг:', str(place.rating)),
                ('Дата:', place.date),
                ('Координаты:', f"{place.mapgeo[0]}, {place.mapgeo[1]}" if place.mapgeo else 'Не указаны')
            ]
            
            for j, (key, value) in enumerate(place_data):
                place_table.cell(j, 0).text = key
                place_table.cell(j, 1).text = value
            
            # Style the place table header
            for cell in place_table.rows[0].cells:
                cell.paragraphs[0].runs[0].bold = True
                cell.paragraphs[0].runs[0].font.size = Pt(11)
            
            # Add some space between places
            doc.add_paragraph()
    
    # Save to bytes
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue() 